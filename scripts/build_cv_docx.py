from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PDF_PATH = ROOT / "public" / "pdfs" / "cv" / "adam-rosinski-cv.pdf"
TMP_DIR = ROOT / "tmp" / "pdfs"
TEXT_PATH = TMP_DIR / "adam-rosinski-cv-layout.txt"
OUTPUT_PATH = ROOT / "output" / "doc" / "adam-rosinski-cv-native.docx"
PDFTOTEXT = Path(r"C:\Program Files\MiKTeX\miktex\bin\x64\pdftotext.exe")


@dataclass
class Entry:
    year: str | None
    title: str
    details: list[str]


SECTION_TITLES = [
    "Keywords",
    "Education",
    "Employment History",
    "Professional Memberships",
    "Monographs",
    "Journal Articles",
    "Book Series Editor",
    "Book Chapters",
    "Journal Reviewing",
    "Conference Presentations",
    "Service on Scientific Committees",
    "Co-supervisor of a doctoral dissertation",
    "Master’s Thesis Supervisor",
    "Sound Engineering",
]


MANUAL_KEYWORDS = (
    "Musical acoustics, electroacoustics, sound directing, sound engineering, music production, "
    "digital and electronic composition and instrumentation, computer music composition, digital "
    "audio workstations (DAW), audiovisual sound design, dubbing, digital technologies in music "
    "education, auditory image analysis."
)


MANUAL_EDUCATION = [
    Entry(
        year="2016",
        title="Doctor of Arts (PhD) in Musical Arts",
        details=[
            "Artistic discipline: Composition and Music Theory.",
            "Awarded on the basis of the doctoral dissertation entitled Combining Sounds into Perceptual Streams: A Comparative Study of Musicians and Non-Musicians.",
            "Degree awarded by resolution of the Council of the Faculty of Composition, Conducting and Music Theory at the Fryderyk Chopin University of Music on 1 July 2016.",
        ],
    ),
    Entry(
        year="2014–2016",
        title="Interfaculty Full-Time Doctoral Programme (Third Cycle)",
        details=["Fryderyk Chopin University of Music."],
    ),
    Entry(
        year="2011",
        title="Master of Arts in Artistic Education in the Field of Musical Arts",
        details=[
            "Specialisation: Music Education.",
            "Awarded on the basis of the master’s thesis entitled Technical Aspects of the Use of Computers in Music Production: Applications of Computer Technology in the Implementation of Selected Didactic Tasks of the Music Teacher.",
            "Faculty of Choral Conducting, Music Education and Eurhythmics, Stanisław Moniuszko Academy of Music in Gdańsk.",
        ],
    ),
]


MANUAL_EMPLOYMENT = [
    Entry(
        year="10.2016 – Present",
        title="University of Warmia and Mazury in Olsztyn",
        details=[
            "Institute of Music, Faculty of Arts.",
            "Assistant Professor.",
        ],
    ),
    Entry(
        year="01.2021 – 06.2024",
        title="University of Warmia and Mazury in Olsztyn",
        details=[
            "Centre for Artistic Initiatives.",
            "Head of the Recording Studio.",
        ],
    ),
    Entry(
        year="10.2017 – 06.2018",
        title="Fryderyk Chopin University of Music",
        details=[
            "Teaching Activities.",
            "Doctoral Seminar.",
        ],
    ),
    Entry(
        year="11.2016 – 12.2016",
        title="University of Gdańsk",
        details=[
            "Faculty of Philology.",
            "Teaching Activities.",
            "Fundamentals of Film Sound Design.",
        ],
    ),
    Entry(
        year="10.2011 – 06.2015",
        title="Academy of Fine Arts in Gdańsk",
        details=[
            "Interfaculty Institute of Art Studies.",
            "Teaching Activities with First- and Second-Cycle Students.",
        ],
    ),
]


MANUAL_MEMBERSHIPS = [
    Entry(
        year="Since 2024",
        title="Professional memberships",
        details=[
            "• American Auditory Society",
            "• Acoustical Society of America",
            "• European Acoustics Association",
            "• Society for Music Production Research",
            "• Society for Music Theory",
            "• Society for Music Analysis",
            "• Society for Music Perception and Cognition",
            "• Cognitive Science Society",
            "• Institute of Electrical and Electronics Engineers",
            "• Association for Computing Machinery",
        ],
    ),
    Entry(
        year="Since 2013",
        title="Professional memberships",
        details=["• Audio Engineering Society"],
    ),
    Entry(
        year="Since 2012",
        title="Professional memberships",
        details=[
            "• Polish Acoustical Society",
            "• Federation of Scientific and Technical Associations – Supreme Technical Organization",
            "• Polish Association of Sound Engineers",
        ],
    ),
]


def fix_text(text: str) -> str:
    if any(token in text for token in ["Ĺ", "Ă", "â", "Â"]):
        try:
            text = text.encode("latin1").decode("utf-8")
        except Exception:
            pass

    replacements = {
        "Rosiñski": "Rosiński",
        "Stanis³aw": "Stanisław",
        "Gdañsk": "Gdańsk",
        "KrakĂłw": "Kraków",
        "Niepe": "Niepe",
        "â€™": "’",
        "â€“": "–",
        "â€”": "—",
        "â€ś": "“",
        "â€ť": "”",
        "â€ž": "„",
        "â€˘": "•",
        "Â·": "·",
        "MiÄ™dzynarodowe": "Międzynarodowe",
        "NiepeĹ‚nosprawnoĹ›Ä‡": "Niepełnosprawność",
        "PrzeglÄ…d": "Przegląd",
        "Kêtrzyńskiego": "Kętrzyńskiego",
        "Mołdoch-Mendoń": "Mołdoch-Mendoń",
        "IsaĂ­as": "Isaías",
        "Sustainbility": "Sustainability",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    translate_map = str.maketrans(
        {
            "ñ": "ń",
            "Ñ": "Ń",
            "³": "ł",
            "£": "Ł",
            "œ": "ś",
            "Œ": "Ś",
            "¹": "ą",
            "¥": "Ą",
            "æ": "ć",
            "Æ": "Ć",
            "ê": "ę",
            "Ê": "Ę",
            "¿": "ż",
            "¯": "Ż",
            "Ÿ": "ź",
        }
    )
    return text.translate(translate_map)


def run_pdftotext() -> list[str]:
    TMP_DIR.mkdir(parents=True, exist_ok=True)
    subprocess.run([str(PDFTOTEXT), "-layout", str(PDF_PATH), str(TEXT_PATH)], check=True)
    text = fix_text(TEXT_PATH.read_text(encoding="utf-8", errors="ignore"))
    pages = [page for page in text.split("\f") if page.strip()]
    cleaned_pages = []
    for index, page in enumerate(pages, start=1):
        lines = []
        for line in page.splitlines():
            stripped = line.rstrip()
            if stripped.strip() == str(index):
                continue
            lines.append(stripped)
        cleaned_pages.append("\n".join(lines))
    return cleaned_pages


def iter_lines(pages: list[str], start: int, end: int) -> list[str]:
    joined = "\n".join(pages[start - 1 : end])
    lines = [line.rstrip() for line in joined.splitlines()]
    return [line for line in lines if line.strip()]


def clean_line(line: str) -> str:
    line = re.sub(r"\s+Download\s*$", "", line)
    line = re.sub(r"\s+the PDF\s*$", "", line)
    return re.sub(r"\s{2,}", " ", line).strip()


def lines_between(lines: list[str], start_marker: str | None = None, end_marker: str | None = None) -> list[str]:
    start_index = 0
    end_index = len(lines)

    if start_marker:
        for idx, line in enumerate(lines):
            if start_marker in line:
                start_index = idx
                break
    if end_marker:
        for idx in range(start_index + 1, len(lines)):
            if end_marker in lines[idx]:
                end_index = idx
                break

    result = []
    for line in lines[start_index:end_index]:
        cleaned = clean_line(line)
        if cleaned:
            result.append(cleaned)
    return result


def find_index(lines: list[str], needle: str) -> int:
    for idx, line in enumerate(lines):
        if needle in line:
            return idx
    return -1


def split_sections(pages: list[str]) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {title: [] for title in SECTION_TITLES}

    page1 = iter_lines(pages, 1, 1)
    page2 = iter_lines(pages, 2, 2)
    page3_4 = iter_lines(pages, 3, 4)
    page4_9 = iter_lines(pages, 4, 9)
    page9 = iter_lines(pages, 9, 9)
    page10_14 = iter_lines(pages, 10, 14)
    page14_24 = iter_lines(pages, 14, 24)
    page24_27 = iter_lines(pages, 24, 27)
    page27_28 = iter_lines(pages, 27, 28)
    page28 = iter_lines(pages, 28, 28)
    page28_30 = iter_lines(pages, 28, 30)

    sections["Keywords"] = lines_between(page1, "Keywords")
    if sections["Keywords"] and sections["Keywords"][0].startswith("Keywords "):
        sections["Keywords"][0] = sections["Keywords"][0].removeprefix("Keywords ").strip()
    sections["Keywords"] = lines_between(sections["Keywords"], None, "Education")

    sections["Education"] = lines_between(page1, "Education")
    if sections["Education"] and sections["Education"][0].startswith("Education "):
        sections["Education"][0] = sections["Education"][0].removeprefix("Education ").strip()

    employment_start = find_index(page2, "Employment History")
    professional_start = find_index(page2, "Professional")
    employment_lines = page2[employment_start:professional_start]
    membership_lines = page2[professional_start:]
    sections["Employment History"] = [clean_line(line) for line in employment_lines if clean_line(line)]
    if sections["Employment History"] and sections["Employment History"][0].startswith("Employment History "):
        sections["Employment History"][0] = sections["Employment History"][0].removeprefix("Employment History ").strip()

    sections["Professional Memberships"] = [clean_line(line) for line in membership_lines if clean_line(line)]
    if sections["Professional Memberships"]:
        sections["Professional Memberships"][0] = "Since 2024"
        sections["Professional Memberships"] = [
            line.replace("Memberships ", "").strip() for line in sections["Professional Memberships"]
        ]

    sections["Monographs"] = lines_between(page3_4, "Monographs", "Journal Articles")
    if sections["Monographs"] and sections["Monographs"][0].startswith("Monographs "):
        sections["Monographs"][0] = sections["Monographs"][0].removeprefix("Monographs ").strip()

    sections["Journal Articles"] = lines_between(page4_9, "Journal Articles", "Book Series Editor")
    if sections["Journal Articles"] and sections["Journal Articles"][0].startswith("Journal Articles "):
        sections["Journal Articles"][0] = sections["Journal Articles"][0].removeprefix("Journal Articles ").strip()

    sections["Book Series Editor"] = lines_between(page9, "Book Series Editor", "Book Chapters")
    if sections["Book Series Editor"] and sections["Book Series Editor"][0].startswith("Book Series Editor "):
        sections["Book Series Editor"][0] = sections["Book Series Editor"][0].removeprefix("Book Series Editor ").strip()

    sections["Book Chapters"] = lines_between(page10_14, "Book Chapters", "Journal Reviewing")
    if sections["Book Chapters"] and sections["Book Chapters"][0].startswith("Book Chapters "):
        sections["Book Chapters"][0] = sections["Book Chapters"][0].removeprefix("Book Chapters ").strip()

    sections["Journal Reviewing"] = lines_between(page14_24, "Journal Reviewing", "Conference Presentations")
    if sections["Journal Reviewing"] and sections["Journal Reviewing"][0].startswith("Journal Reviewing "):
        sections["Journal Reviewing"][0] = sections["Journal Reviewing"][0].removeprefix("Journal Reviewing ").strip()

    conference_start = find_index(page24_27, "Conference")
    service_start = find_index(page27_28, "Service on Scientific")
    co_supervisor_start = find_index(page28, "Co-supervisor")
    co_supervisor_start_combined = find_index(page27_28, "Co-supervisor")
    masters_start = find_index(page28, "Master’s Thesis")
    masters_start_combined = find_index(page27_28, "Master’s Thesis")
    sound_start = find_index(page28_30, "Sound Engineering")

    conference_lines = page24_27[conference_start: find_index(page24_27, "Service on Scientific")] if find_index(page24_27, "Service on Scientific") != -1 else page24_27[conference_start:]
    sections["Conference Presentations"] = [clean_line(line) for line in conference_lines if clean_line(line)]
    if sections["Conference Presentations"][:2] == ["Conference 2025", "Presentations"]:
        sections["Conference Presentations"] = ["2025"] + sections["Conference Presentations"][2:]

    service_lines = page27_28[service_start:co_supervisor_start_combined]
    sections["Service on Scientific Committees"] = [clean_line(line) for line in service_lines if clean_line(line)]
    if sections["Service on Scientific Committees"]:
        sections["Service on Scientific Committees"][0] = (
            sections["Service on Scientific Committees"][0].replace("Service on Scientific", "").strip() or "2025"
        )
        sections["Service on Scientific Committees"] = [
            line for line in sections["Service on Scientific Committees"] if line != "Committees"
        ]

    co_lines = page28[co_supervisor_start:masters_start]
    sections["Co-supervisor of a doctoral dissertation"] = [clean_line(line) for line in co_lines if clean_line(line)]
    if sections["Co-supervisor of a doctoral dissertation"]:
        sections["Co-supervisor of a doctoral dissertation"][0] = sections["Co-supervisor of a doctoral dissertation"][0].replace("Co-supervisor ", "").strip()
        if len(sections["Co-supervisor of a doctoral dissertation"]) > 2:
            sections["Co-supervisor of a doctoral dissertation"][2] = (
                sections["Co-supervisor of a doctoral dissertation"][2].replace("of a doctoral ", "").strip()
            )
        if len(sections["Co-supervisor of a doctoral dissertation"]) > 3:
            sections["Co-supervisor of a doctoral dissertation"][3] = (
                sections["Co-supervisor of a doctoral dissertation"][3].replace("dissertation ", "").strip()
            )

    master_lines = page28[masters_start:sound_start]
    sections["Master’s Thesis Supervisor"] = [clean_line(line) for line in master_lines if clean_line(line)]
    if sections["Master’s Thesis Supervisor"]:
        sections["Master’s Thesis Supervisor"][0] = sections["Master’s Thesis Supervisor"][0].replace(
            "Master’s Thesis ", ""
        ).strip()
        if len(sections["Master’s Thesis Supervisor"]) > 2:
            sections["Master’s Thesis Supervisor"][2] = sections["Master’s Thesis Supervisor"][2].replace("Supervisor ", "").strip()

    sections["Sound Engineering"] = [clean_line(line) for line in page28_30[sound_start:] if clean_line(line)]
    if sections["Sound Engineering"]:
        sections["Sound Engineering"][0] = sections["Sound Engineering"][0].replace("Sound Engineering ", "").strip()

    return sections


def parse_keywords(lines: list[str]) -> str:
    return " ".join(line for line in lines if line).strip()


def parse_education(lines: list[str]) -> list[Entry]:
    entries: list[Entry] = []
    current_year = None
    current_details: list[str] = []
    for line in lines:
        if re.fullmatch(r"\d{4}(?:\s*-\s*\d{4})?", line):
            if current_year or current_details:
                entries.append(Entry(year=current_year, title=current_details[0], details=current_details[1:]))
            current_year = line
            current_details = []
            continue
        if line:
            current_details.append(line)
    if current_year or current_details:
        entries.append(Entry(year=current_year, title=current_details[0], details=current_details[1:]))
    return entries


def parse_employment(lines: list[str]) -> list[Entry]:
    entries: list[Entry] = []
    current_org: list[str] = []
    current_details: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if re.fullmatch(r"\d{2}\.\d{4}\s*-\s*(?:Present|\d{2}\.\d{4})", line):
            if current_org or current_details:
                entries.append(
                    Entry(
                        year=current_details[0] if current_details else None,
                        title=" / ".join(current_org) if current_org else "",
                        details=current_details[1:],
                    )
                )
                current_details = []
            if not current_org:
                current_org = ["Employment item"]
            current_details = [line]
            if i + 1 < len(lines):
                current_details.append(lines[i + 1])
                i += 1
            if i + 1 < len(lines):
                nxt = lines[i + 1]
                if not re.fullmatch(r"\d{2}\.\d{4}\s*-\s*(?:Present|\d{2}\.\d{4})", nxt):
                    current_details.append(nxt)
                    i += 1
            entries.append(
                Entry(
                    year=current_details[0],
                    title=" / ".join(current_org),
                    details=current_details[1:],
                )
            )
            current_details = []
        else:
            current_org.append(line)
        i += 1
    return [entry for entry in entries if entry.title]


def parse_memberships(lines: list[str]) -> list[Entry]:
    entries: list[Entry] = []
    current_year = None
    current_items: list[str] = []
    for line in lines:
        if line.startswith("Since "):
            if current_year or current_items:
                entries.append(Entry(year=current_year, title=current_year or "", details=current_items))
            current_year = line
            current_items = []
            continue
        if line.startswith("• "):
            current_items.append(line.removeprefix("• ").strip())
            continue
        if current_items and not line.startswith("Since "):
            current_items[-1] = f"{current_items[-1]} {line}".strip()
    if current_year or current_items:
        entries.append(Entry(year=current_year, title=current_year or "", details=current_items))
    return entries


def parse_author_entries(lines: list[str], author_markers: tuple[str, ...] = ("A. Rosiński", "M. Waśkiewicz, A. Rosiński")) -> list[Entry]:
    entries: list[Entry] = []
    current_year = None
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        if not current_lines:
            return
        while current_lines and current_lines[-1] == "":
            current_lines.pop()
        if not current_lines:
            return
        title_parts: list[str] = []
        detail_parts: list[str] = []
        mode = "title"
        for part in current_lines:
            if part.startswith(author_markers):
                continue
            if mode == "title" and (
                "”" in part
                or "„" in part
                or "Press" in part
                or part.startswith("University")
                or part.startswith("Przestrzenie")
                or part.startswith("Sonic ")
                or part.startswith("20th ")
                or part.startswith("II Sobór")
                or part.startswith("Nowe trendy")
                or part.startswith("Młodzi naukowcy")
                or part.startswith("Muzyka sakralna")
                or part.startswith("Proceedings")
                or part.startswith("Humanitarian Corpus")
            ):
                mode = "details"
            if mode == "title":
                title_parts.append(part)
            else:
                detail_parts.append(part)
        if not title_parts and current_lines:
            title_parts = [current_lines[0]]
            detail_parts = current_lines[1:]
        title = " ".join(title_parts).strip().rstrip(",")
        entries.append(Entry(year=current_year, title=title, details=detail_parts))
        current_lines = []

    for line in lines:
        if re.fullmatch(r"\d{4}", line):
            flush()
            current_year = line
            continue
        if line in author_markers:
            flush()
            current_lines = [line]
            continue
        if line.endswith("(ed.)") or line.endswith("(red.)"):
            flush()
            current_lines = [line]
            continue
        if current_lines:
            current_lines.append(line)
    flush()
    return entries


def parse_reviewing(lines: list[str]) -> list[Entry]:
    entries: list[Entry] = []
    current_year = None
    current_title: list[str] = []
    current_details: list[str] = []

    def flush() -> None:
        nonlocal current_title, current_details
        if current_title:
            entries.append(
                Entry(
                    year=current_year,
                    title=" ".join(current_title).strip(),
                    details=[detail for detail in current_details if detail],
                )
            )
        current_title = []
        current_details = []

    for line in lines:
        if re.fullmatch(r"\d{4}", line):
            flush()
            current_year = line
            continue
        if line.startswith(("“", "„")):
            current_details.append(line)
            continue
        if line.startswith(("ISSN:", "e-ISSN:")):
            current_details.append(line)
            flush()
            continue
        if not current_details:
            current_title.append(line)
        else:
            current_details.append(line)
    flush()
    return entries


def parse_event_entries(lines: list[str], author_first: bool = True) -> list[Entry]:
    entries: list[Entry] = []
    current_year = None
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_lines
        if not current_lines:
            return
        lines_local = [line for line in current_lines if line]
        if author_first and lines_local and lines_local[0] == "A. Rosiński":
            lines_local = lines_local[1:]
        title_parts: list[str] = []
        detail_parts: list[str] = []
        for idx, part in enumerate(lines_local):
            if idx > 0 and (
                re.search(r"\bConference\b", part)
                or "Foundation" in part
                or "University" in part
                or "Lublin/online" in part
                or re.search(r"\d{1,2}\s", part)
            ):
                detail_parts = lines_local[idx:]
                break
            title_parts.append(part)
        if not detail_parts and len(lines_local) > len(title_parts):
            detail_parts = lines_local[len(title_parts) :]
        entries.append(Entry(year=current_year, title=" ".join(title_parts).strip(), details=detail_parts))
        current_lines = []

    for line in lines:
        if re.fullmatch(r"\d{4}", line):
            flush()
            current_year = line
            continue
        if author_first and line == "A. Rosiński":
            flush()
            current_lines = [line]
            continue
        if not author_first and not current_lines:
            current_lines = [line]
            continue
        if current_lines and (
            (author_first and line == "A. Rosiński")
            or (not author_first and re.search(r"\d{1,2}\s(?:January|February|March|April|May|June|July|August|September|October|November|December)", current_lines[-1]))
        ):
            flush()
            current_lines = [line]
            continue
        if current_lines:
            current_lines.append(line)
        else:
            current_lines = [line]
        if not author_first and re.search(r"\d{1,2}\s(?:January|February|March|April|May|June|July|August|September|October|November|December)", line):
            flush()
    flush()
    return entries


def parse_supervision(lines: list[str]) -> list[Entry]:
    entries: list[Entry] = []
    chunks: list[list[str]] = []
    current: list[str] = []
    for line in lines:
        if line.startswith("Date of Defence:") and current:
            current.append(line)
            chunks.append(current)
            current = []
            continue
        current.append(line)
    if current:
        chunks.append(current)
    for chunk in chunks:
        title_parts = chunk[:3]
        details = chunk[3:]
        if title_parts and title_parts[-1].endswith("(Written in") and details:
            title_parts[-1] = f"{title_parts[-1]} {details[0]}".strip()
            details = details[1:]
        entries.append(Entry(year=None, title=" ".join(title_parts).strip(), details=details))
    return entries


def parse_sound_engineering(lines: list[str]) -> list[Entry]:
    entries: list[Entry] = []
    chunks: list[list[str]] = []
    current: list[str] = []
    for line in lines:
        if line.startswith("Running time:") and current:
            current.append(line)
            chunks.append(current)
            current = []
            continue
        current.append(line)
    if current:
        chunks.append(current)
    for chunk in chunks:
        title_parts: list[str] = []
        details: list[str] = []
        for idx, line in enumerate(chunk):
            if idx > 0 and (
                line.startswith("Director:")
                or line.startswith("Performance:")
                or line.startswith("Audio Preparation")
                or line.startswith("Żurawscy")
                or line.startswith("Brąswałd")
                or line.startswith("Biskupia")
                or line.startswith("Sprawa")
                or line.startswith("Production:")
                or re.fullmatch(r"\d{4}", line)
            ):
                details = chunk[idx:]
                break
            title_parts.append(line)
        if not details and len(chunk) > len(title_parts):
            details = chunk[len(title_parts) :]
        entries.append(Entry(year=None, title=" ".join(title_parts).strip(), details=details))
    return entries


def ensure_styles(doc: Document) -> None:
    styles = doc.styles

    if "CV Section Heading" not in styles:
        style = styles.add_style("CV Section Heading", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Heading 1"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor(24, 68, 104)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)

    if "CV Entry Title" not in styles:
        style = styles.add_style("CV Entry Title", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(1)

    if "CV Entry Detail" not in styles:
        style = styles.add_style("CV Entry Detail", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.2)
        style.paragraph_format.space_after = Pt(1)
        style.paragraph_format.line_spacing = 1.08

    if "CV Meta" not in styles:
        style = styles.add_style("CV Meta", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(9.5)
        style.font.italic = True
        style.font.color.rgb = RGBColor(90, 98, 108)
        style.paragraph_format.space_after = Pt(2)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def add_header(doc: Document, keywords: str) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Inches(4.75)
    right.width = Inches(2.1)
    for cell in (left, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_border(
            cell,
            top={"val": "nil"},
            left={"val": "nil"},
            bottom={"val": "nil"},
            right={"val": "nil"},
        )

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("dr Adam Rosiński")
    r.font.name = "Aptos Display"
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(24, 68, 104)

    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(3)
    run = p2.add_run("Curriculum Vitae")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True

    p3 = left.add_paragraph()
    p3.paragraph_format.space_before = Pt(2)
    run = p3.add_run("Contact: adamrosinski.com/contact")
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(56, 64, 72)

    rp = right.paragraphs[0]
    rp.style = doc.styles["CV Meta"]
    rr = rp.add_run("Keywords")
    rr.bold = True
    rr.font.italic = False
    rp2 = right.add_paragraph(keywords)
    rp2.style = doc.styles["CV Entry Detail"]
    rp2.paragraph_format.left_indent = Inches(0)

    divider = doc.add_paragraph()
    divider.paragraph_format.space_before = Pt(8)
    divider.paragraph_format.space_after = Pt(4)
    run = divider.add_run()
    run.add_break(WD_BREAK.LINE)
    ppr = divider._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "8FA7BC")
    p_bdr.append(bottom)
    ppr.append(p_bdr)


def add_section_heading(doc: Document, title: str) -> None:
    doc.add_paragraph(title, style="CV Section Heading")


def add_entry(doc: Document, entry: Entry, year_inline: bool = True) -> None:
    if entry.year:
        p = doc.add_paragraph(style="CV Meta")
        p.add_run(entry.year)
    p = doc.add_paragraph(style="CV Entry Title")
    p.add_run(entry.title)
    for detail in entry.details:
        if not detail:
            continue
        d = doc.add_paragraph(style="CV Entry Detail")
        if detail.startswith("• "):
            d.style = doc.styles["List Bullet"]
            d.add_run(detail.removeprefix("• ").strip())
        else:
            d.add_run(detail)


def build_doc() -> Path:
    pages = run_pdftotext()
    sections = split_sections(pages)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    ensure_styles(doc)

    add_header(doc, MANUAL_KEYWORDS)

    add_section_heading(doc, "Education")
    for entry in MANUAL_EDUCATION:
        add_entry(doc, entry)

    add_section_heading(doc, "Employment History")
    for entry in MANUAL_EMPLOYMENT:
        add_entry(doc, entry)

    add_section_heading(doc, "Professional Memberships")
    for entry in MANUAL_MEMBERSHIPS:
        add_entry(doc, entry)

    for title in ("Monographs", "Journal Articles", "Book Series Editor", "Book Chapters"):
        add_section_heading(doc, title)
        for entry in parse_author_entries(sections[title]):
            add_entry(doc, entry)

    add_section_heading(doc, "Journal Reviewing")
    for entry in parse_reviewing(sections["Journal Reviewing"]):
        add_entry(doc, entry)

    add_section_heading(doc, "Conference Presentations")
    for entry in parse_event_entries(sections["Conference Presentations"], author_first=True):
        add_entry(doc, entry)

    add_section_heading(doc, "Service on Scientific Committees")
    for entry in parse_event_entries(sections["Service on Scientific Committees"], author_first=False):
        add_entry(doc, entry)

    add_section_heading(doc, "Co-supervisor of a doctoral dissertation")
    for entry in parse_supervision(sections["Co-supervisor of a doctoral dissertation"]):
        add_entry(doc, entry)

    add_section_heading(doc, "Master’s Thesis Supervisor")
    for entry in parse_supervision(sections["Master’s Thesis Supervisor"]):
        add_entry(doc, entry)

    add_section_heading(doc, "Sound Engineering")
    for entry in parse_sound_engineering(sections["Sound Engineering"]):
        add_entry(doc, entry)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Adam Rosiński CV"
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.subject = "Native editable CV document rebuilt from PDF source"
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
